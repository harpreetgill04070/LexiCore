# # # vector_database.py

# # # vector_database.py
# import os
# import concurrent.futures
# from dotenv import load_dotenv
# from langchain_community.document_loaders import PDFPlumberLoader
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_openai import OpenAIEmbeddings
# from langchain_pinecone import PineconeVectorStore
# from pinecone import Pinecone, ServerlessSpec

# load_dotenv()

# PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# INDEX_NAME = "lexicore"
# pdfs_directory = "pdfs/"

# # 🔧 Initialize Pinecone
# pc = Pinecone(api_key=PINECONE_API_KEY)

# # 🔧 Setup OpenAI Embedding Model
# embedding_model = OpenAIEmbeddings(
#     model="text-embedding-3-small",  # cheap + fast, 1536 dimensions
#     api_key=OPENAI_API_KEY
# )
# EMBEDDING_DIM = 1536

# # 🔧 Ensure Pinecone index exists
# existing_indexes = [i["name"] for i in pc.list_indexes()]
# if INDEX_NAME not in existing_indexes:
#     print(f"📌 Creating Pinecone index: {INDEX_NAME}")
#     pc.create_index(
#         name=INDEX_NAME,
#         dimension=EMBEDDING_DIM,
#         metric="cosine",
#         spec=ServerlessSpec(cloud="aws", region="us-east-1")
#     )
# else:
#     print(f"✅ Pinecone index '{INDEX_NAME}' already exists.")

# # ---- PDF & Chunk Handling ----
# def upload_pdf(file):
#     os.makedirs(pdfs_directory, exist_ok=True)
#     with open(os.path.join(pdfs_directory, file.name), "wb") as f:
#         f.write(file.getbuffer())

# def load_pdf(file_path):
#     loader = PDFPlumberLoader(file_path)
#     return loader.load()

# def create_chunks(documents):
#     splitter = RecursiveCharacterTextSplitter(
#         chunk_size=800,
#         chunk_overlap=100,
#         add_start_index=True
#     )
#     return splitter.split_documents(documents)

# # ---- Embedding + Upsert ----
# def store_in_pinecone(documents):
#     print(f"🔄 Generating embeddings for {len(documents)} chunks...")
#     texts = [doc.page_content for doc in documents]

#     with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
#         embeddings = list(executor.map(embedding_model.embed_query, texts))

#     vectors = [
#         {"id": f"doc-{i}", "values": emb, "metadata": {"text": doc.page_content}}
#         for i, (doc, emb) in enumerate(zip(documents, embeddings))
#     ]

#     print(f"🚀 Uploading {len(vectors)} vectors to Pinecone...")
#     index = pc.Index(INDEX_NAME)

#     batch_size = 100
#     for i in range(0, len(vectors), batch_size):
#         index.upsert(vectors=vectors[i:i+batch_size])

#     print("✅ All vectors uploaded to Pinecone.")

# def get_vectorstore():
#     return PineconeVectorStore(index_name=INDEX_NAME, embedding=embedding_model)

# -----------------------------------------------------------------------------------------------
import os
import uuid
import time
import concurrent.futures
from typing import List
from dotenv import load_dotenv
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone, ServerlessSpec
from langchain.schema import Document
from docx import Document as DocxDocument
import pandas as pd
from zipfile import ZipFile
import re  # Added for regex splitting

load_dotenv()

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

INDEX_NAME = "ross-ai-index"
uploads_dir = "pdfs"   # keep existing path, but we’ll store all file types here
os.makedirs(uploads_dir, exist_ok=True)

# ---- Pinecone + Embeddings ----
pc = Pinecone(api_key=PINECONE_API_KEY)

embedding_model = OpenAIEmbeddings(
    model="text-embedding-3-small",
    api_key=OPENAI_API_KEY,
)
EMBEDDING_DIM = 1536

existing = [i["name"] for i in pc.list_indexes()]
if INDEX_NAME not in existing:
    print(f"📌 Creating Pinecone index: {INDEX_NAME}")
    pc.create_index(
        name=INDEX_NAME,
        dimension=EMBEDDING_DIM,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
    )
    # Wait for index to be ready (serverless can take 1-2 min)
    print("⏳ Waiting for index to be fully ready...")
    while not pc.describe_index(INDEX_NAME).status.ready:
        time.sleep(10)
    print("✅ Index is ready.")
else:
    print(f"✅ Pinecone index '{INDEX_NAME}' already exists.")

# ---- Helpers: loaders for multiple types ----
def _load_pdf(path: str) -> List[Document]:
    loader = PDFPlumberLoader(path)
    pages = loader.load()
    # Concat all page content
    full_text = ""
    page_map = {}  # To track page for each part
    current_page = 1
    for page_doc in pages:
        full_text += page_doc.page_content + "\n\n"
        # Track approximate position
        page_map[len(full_text)] = str(page_doc.metadata.get("page", current_page))
        current_page += 1
    
    # Improved: Split on 'Article X' or 'Preamble' for structured docs
    # Find sections
    sections = []
    matches = list(re.finditer(r'(Preamble|Article \d+)', full_text))
    if matches:
        start = 0
        for match in matches:
            end = match.start()
            if start != end:
                section_content = full_text[start:end].strip()
                if section_content:
                    # Find approximate page
                    pos = start
                    page = next((p for l, p in page_map.items() if l > pos), str(current_page))
                    md = {
                        "source": os.path.basename(path),
                        "filetype": "pdf",
                        "path": path,
                        "page": page,
                        "section": full_text[start-20:start].strip() if start > 0 else "Preamble"  # Previous section name
                    }
                    sections.append(Document(page_content=section_content, metadata=md))
            start = end
        # Last section
        section_content = full_text[start:].strip()
        if section_content:
            pos = start
            page = next((p for l, p in page_map.items() if l > pos), str(current_page))
            md = {
                "source": os.path.basename(path),
                "filetype": "pdf",
                "path": path,
                "page": page,
                "section": full_text[start:start+20].strip()
            }
            sections.append(Document(page_content=section_content, metadata=md))
    else:
        # Fallback to per-page if no sections found
        sections = pages
        for d in sections:
            d.metadata["source"] = os.path.basename(d.metadata.get("source", path))
            d.metadata["filetype"] = "pdf"
            d.metadata["path"] = path
            if "page" in d.metadata and isinstance(d.metadata["page"], int):
                d.metadata["page"] = str(d.metadata["page"])
    
    return sections

def _load_docx(path: str) -> List[Document]:
    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    md = {"source": os.path.basename(path), "filetype": "docx", "path": path}
    return [Document(page_content=text, metadata=md)]

def _load_xlsx(path: str) -> List[Document]:
    xl = pd.ExcelFile(path)
    docs = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        text = df.to_csv(index=False)
        md = {
            "source": os.path.basename(path),
            "filetype": "xlsx",
            "sheet": sheet,
            "path": path,
        }
        docs.append(Document(page_content=text, metadata=md))
    return docs

def _load_zip(path: str) -> List[Document]:
    out = []
    extract_dir = os.path.join(uploads_dir, f"unzipped_{uuid.uuid4().hex[:8]}")
    os.makedirs(extract_dir, exist_ok=True)
    with ZipFile(path, "r") as z:
        z.extractall(extract_dir)
    for root, _, files in os.walk(extract_dir):
        for name in files:
            out.extend(load_any(os.path.join(root, name)))
    return out

def load_any(path: str) -> List[Document]:
    low = path.lower()
    if low.endswith(".pdf"):
        return _load_pdf(path)
    if low.endswith(".docx"):
        return _load_docx(path)
    if low.endswith(".xlsx"):
        return _load_xlsx(path)
    if low.endswith(".zip"):
        return _load_zip(path)
    # ignore other types
    return []

def create_chunks(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=120,
        add_start_index=True,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    # preserve key metadata explicitly
    for i, c in enumerate(chunks):
        c.metadata["chunk_id"] = i
        c.metadata["source"] = os.path.basename(c.metadata.get("source", ""))
    return chunks

def store_in_pinecone(documents: List[Document]) -> None:
    if not documents:
        return
    print(f"🔄 Generating embeddings for {len(documents)} chunks...")
    texts = [d.page_content for d in documents]
    # batch embed (keeps order)
    embeddings = embedding_model.embed_documents(texts)

    # compose metadata (text + useful fields) - Fix: Ensure all values are valid types (str, float, etc.)
    vectors = []
    index = pc.Index(INDEX_NAME)
    for d, emb in zip(documents, embeddings):
        meta = {
            "text": d.page_content,  # str
            "source": str(d.metadata.get("source", "")),  # Ensure str
            "page": str(d.metadata.get("page", "")) if d.metadata.get("page") is not None else None,  # str or None (None is okay, skipped)
            "sheet": str(d.metadata.get("sheet", "")) if d.metadata.get("sheet") is not None else None,
            "filetype": str(d.metadata.get("filetype", "")),
            "path": str(d.metadata.get("path", "")),
            "chunk_id": str(d.metadata.get("chunk_id", "")),
            "section": str(d.metadata.get("section", ""))  # Added for article/section
        }
        # Remove None values (Pinecone skips them)
        meta = {k: v for k, v in meta.items() if v is not None}
        vec_id = str(uuid.uuid4().hex)  # Ensure str ID
        vectors.append(
            {
                "id": vec_id,
                "values": [float(val) for val in emb],  # Ensure list of floats
                "metadata": meta,
            }
        )

    print(f"🚀 Uploading {len(vectors)} vectors to Pinecone...")
    batch = 100
    total_upserted = 0
    for i in range(0, len(vectors), batch):
        batch_vectors = vectors[i : i + batch]
        try:
            result = index.upsert(vectors=batch_vectors)
            total_upserted += len(batch_vectors)
            # Verify after each batch (eventual consistency: wait 5s and check stats)
            time.sleep(5)
            stats = index.describe_index_stats()
            current_count = stats.total_vector_count
            print(f"📊 After batch {i//batch + 1}: Upserted {len(batch_vectors)}, Total vectors in index: {current_count}")
            if current_count < total_upserted:
                print("⚠️ Warning: Vectors may take time to index due to eventual consistency. Retrying check...")
                time.sleep(10)
                stats = index.describe_index_stats()
                print(f"📊 Re-check: Total vectors in index: {stats.total_vector_count}")
        except Exception as e:
            print(f"❌ Error upserting batch {i//batch + 1}: {e}")
            raise

    # Final verification
    time.sleep(10)  # Allow final consistency
    final_stats = index.describe_index_stats()
    print(f"✅ All {total_upserted} vectors uploaded. Final index stats: {final_stats.total_vector_count} total vectors.")

def get_vectorstore() -> PineconeVectorStore:
    return PineconeVectorStore(index_name=INDEX_NAME, embedding=embedding_model)

def get_retriever(k: int = 20, file_filter: str | None = None):  # Increased default k to 20 for better coverage of documents
    store = get_vectorstore()
    search_kwargs = {"k": k}
    filter_ = None
    if file_filter:
        # exact match on filename stored as metadata 'source'
        filter_ = {"source": {"$eq": os.path.basename(file_filter)}}
    return store.as_retriever(search_kwargs=search_kwargs, filter=filter_)